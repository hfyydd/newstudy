"""
PDF文本提取测试脚本

测试用例：
1. 纯文本PDF提取
2. 图片PDF（扫描件）OCR
3. 混合PDF（文本+图片）
4. 错误处理
"""

import io
import sys
from pathlib import Path

# 添加当前目录到Python路径
sys.path.insert(0, str(Path(__file__).parent))

from file_text_extractor import extract_text_from_upload


def create_text_pdf_content() -> bytes:
    """创建一个简单的文本PDF（用于测试普通PDF提取）"""
    # PDF内容 - 简单的英文文本
    pdf_content = b"""%PDF-1.4
1 0 obj
<< /Type /Catalog /Pages 2 0 R >>
endobj
2 0 obj
<< /Type /Pages /Kids [3 0 R] /Count 1 >>
endobj
3 0 obj
<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 792] /Contents 4 0 R /Resources << /Font << /F1 5 0 R >> >> >>
endobj
4 0 obj
<< /Length 200 >>
stream
BT
/F1 24 Tf
100 700 Td
(This is a test PDF document) Tj
0 -30 Td
(Containing both Chinese and English text) Tj
0 -30 Td
(For testing PDF text extraction functionality) Tj
ET
endstream
endobj
5 0 obj
<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>
endobj
xref
0 6
0000000000 65535 f 
0000000009 00000 n 
0000000058 00000 n 
0000000115 00000 n 
0000000266 00000 n 
0000000518 00000 n 
trailer
<< /Size 6 /Root 1 0 R >>
startxref
597
%%EOF"""
    return pdf_content


def test_text_pdf_extraction():
    """测试纯文本PDF提取"""
    print("=" * 60)
    print("测试1: 纯文本PDF提取")
    print("=" * 60)
    
    pdf_content = create_text_pdf_content()
    
    try:
        text = extract_text_from_upload("test_text.pdf", pdf_content)
        print(f"提取成功！文本长度: {len(text)} 字符")
        print(f"提取的文本:\n{text}")
        assert len(text) > 0, "PDF文本提取为空"
        assert "test PDF" in text.lower() or "pdf" in text.lower(), f"PDF文本内容不正确: {text}"
        print("✅ 测试通过\n")
        return True
    except Exception as e:
        print(f"❌ 测试失败: {e}\n")
        return False


def test_txt_extraction():
    """测试纯文本文件提取"""
    print("=" * 60)
    print("测试2: 纯文本文件提取")
    print("=" * 60)
    
    txt_content = b"""This is a test document
With multiple lines of text
For testing text extraction function
    
This is an introduction about machine learning:
Machine learning is a branch of AI that enables computers to learn from data without being explicitly programmed.
"""
    
    try:
        text = extract_text_from_upload("test.txt", txt_content)
        print(f"提取成功！文本长度: {len(text)} 字符")
        print(f"提取的文本:\n{text}")
        assert len(text) > 0, "TXT文本提取为空"
        assert "machine learning" in text.lower(), "TXT文本内容不正确"
        print("✅ 测试通过\n")
        return True
    except Exception as e:
        print(f"❌ 测试失败: {e}\n")
        return False


def test_markdown_extraction():
    """测试Markdown文件提取"""
    print("=" * 60)
    print("测试3: Markdown文件提取")
    print("=" * 60)
    
    md_content = b"""# Test Document

## This is a heading

This is **bold** text, this is *italic* text.

- List item 1
- List item 2
- List item 3

```python
def hello():
    print("Hello, World!")
```
"""
    
    try:
        text = extract_text_from_upload("test.md", md_content)
        print(f"提取成功！文本长度: {len(text)} 字符")
        print(f"提取的文本:\n{text}")
        assert len(text) > 0, "Markdown文本提取为空"
        assert "Test Document" in text or "heading" in text.lower(), "Markdown文本内容不正确"
        print("✅ 测试通过\n")
        return True
    except Exception as e:
        print(f"❌ 测试失败: {e}\n")
        return False


def test_docx_extraction():
    """测试Word文档提取"""
    print("=" * 60)
    print("测试4: Word文档提取")
    print("=" * 60)
    
    try:
        from docx import Document
        doc = Document()
        doc.add_heading('Test Document', 0)
        doc.add_paragraph('This is a test Word document')
        doc.add_paragraph('For testing text extraction function')
        
        # 保存到内存
        docx_buffer = io.BytesIO()
        doc.save(docx_buffer)
        docx_content = docx_buffer.getvalue()
        
        text = extract_text_from_upload("test.docx", docx_content)
        print(f"提取成功！文本长度: {len(text)} 字符")
        print(f"提取的文本:\n{text}")
        assert len(text) > 0, "DOCX文本提取为空"
        assert "test" in text.lower(), "DOCX文本内容不正确"
        print("✅ 测试通过\n")
        return True
    except ImportError:
        print("⚠️ 跳过测试: python-docx未安装\n")
        return True
    except Exception as e:
        print(f"❌ 测试失败: {e}\n")
        return False


def test_unsupported_format():
    """测试不支持的文件格式"""
    print("=" * 60)
    print("测试5: 不支持的文件格式")
    print("=" * 60)
    
    content = b"some binary data"
    
    try:
        text = extract_text_from_upload("test.exe", content)
        print(f"❌ 测试失败: 应该抛出异常但没有\n")
        return False
    except ValueError as e:
        if "不支持" in str(e):
            print(f"✅ 测试通过: 正确抛出异常 - {e}\n")
            return True
        else:
            print(f"❌ 测试失败: 异常消息不正确 - {e}\n")
            return False
    except Exception as e:
        print(f"❌ 测试失败: 异常类型不正确 - {e}\n")
        return False


def test_image_ocr():
    """测试图片OCR（如果安装了easyocr）"""
    print("=" * 60)
    print("测试6: 图片OCR")
    print("=" * 60)
    
    try:
        import numpy as np
        import cv2
        
        # 创建一个简单的测试图片（包含文字）
        # 使用OpenCV创建一个白色图片
        img = np.ones((200, 600, 3), dtype=np.uint8) * 255
        
        # 添加中文和英文文字
        import PIL.Image
        from PIL import ImageDraw, ImageFont
        
        pil_img = PIL.Image.fromarray(cv2.cvtColor(img, cv2.COLOR_BGR2RGB))
        draw = ImageDraw.Draw(pil_img)
        
        # 尝试使用默认字体
        try:
            font = ImageFont.truetype("/System/Library/Fonts/PingFang.ttc", 24)
        except:
            font = ImageFont.load_default()
        
        draw.text((20, 30), "Test OCR Function", fill=(0, 0, 0), font=font)
        draw.text((20, 70), "Hello World", fill=(0, 0, 0), font=font)
        
        # 转换回OpenCV格式
        img_array = np.array(pil_img)
        img = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
        
        # 编码为JPEG
        success, encoded_img = cv2.imencode('.jpg', img)
        if not success:
            raise ValueError("无法编码图片")
        
        img_bytes = encoded_img.tobytes()
        
        # 使用函数提取
        text = extract_text_from_upload("test.jpg", img_bytes)
        print(f"OCR提取成功！文本长度: {len(text)} 字符")
        print(f"提取的文本: {text}")
        assert len(text) > 0, "OCR提取为空"
        print("✅ 测试通过\n")
        return True
    except ImportError as e:
        print(f"⚠️ 跳过测试: 缺少依赖 - {e}\n")
        return True
    except Exception as e:
        print(f"❌ 测试失败: {e}\n")
        return False


def run_all_tests():
    """运行所有测试"""
    print("\n" + "=" * 60)
    print("PDF文本提取功能测试")
    print("=" * 60 + "\n")
    
    results = []
    
    # 运行测试
    results.append(("纯文本PDF提取", test_text_pdf_extraction()))
    results.append(("纯文本文件提取", test_txt_extraction()))
    results.append(("Markdown文件提取", test_markdown_extraction()))
    results.append(("Word文档提取", test_docx_extraction()))
    results.append(("不支持的格式", test_unsupported_format()))
    results.append(("图片OCR", test_image_ocr()))
    
    # 输出结果汇总
    print("=" * 60)
    print("测试结果汇总")
    print("=" * 60)
    
    passed = 0
    failed = 0
    
    for name, result in results:
        status = "✅ 通过" if result else "❌ 失败"
        print(f"{name}: {status}")
        if result:
            passed += 1
        else:
            failed += 1
    
    print(f"\n总计: {passed} 通过, {failed} 失败")
    
    return failed == 0


if __name__ == "__main__":
    success = run_all_tests()
    sys.exit(0 if success else 1)
